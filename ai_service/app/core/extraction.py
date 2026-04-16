from __future__ import annotations

import csv
import hashlib
import io
import logging
import re
import unicodedata
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.config import settings
from app.core.ocr.factory import create_ocr_provider

# Suppress noisy pdfminer/pdfplumber font warnings (FontBBox not parseable)
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)

# ---------------------------------------------------------------------------
# Arabic / RTL helpers
# ---------------------------------------------------------------------------

# Unicode ranges that indicate RTL text (Arabic base + Presentation Forms)
_RTL_RANGE_RE = re.compile(
    r"[\u0600-\u06FF"  # Arabic
    r"\u0750-\u077F"  # Arabic Supplement
    r"\u08A0-\u08FF"  # Arabic Extended-A
    r"\uFB50-\uFDFF"  # Arabic Presentation Forms-A
    r"\uFE70-\uFEFF]"  # Arabic Presentation Forms-B
)


def _has_rtl(text: str) -> bool:
    """Return True if *text* contains Arabic/RTL characters."""
    return bool(_RTL_RANGE_RE.search(text))


def _clean_word(word: str) -> str:
    """Strip null bytes and NFKC-normalize a single word token."""
    return unicodedata.normalize("NFKC", word.replace("\x00", ""))


def _fix_rtl_line(line: str) -> str:
    """Strip null bytes then NFKC-normalize Arabic presentation forms.

    Many Arabic PDFs store text with:
      1. Presentation form codepoints (U+FB50–U+FEFF) instead of base chars.
      2. Null bytes (U+0000) interspersed between glyph characters.

    Stripping null bytes first prevents character fragmentation; NFKC then
    converts presentation forms to canonical base Arabic (U+0600–U+06FF).
    """
    line = line.replace("\x00", "")
    return unicodedata.normalize("NFKC", line)


def _fix_rtl_text(text: str) -> str:
    """Apply RTL fix line-by-line only when RTL characters are present."""
    if not _has_rtl(text):
        return text
    lines = text.splitlines()
    fixed = [_fix_rtl_line(line) if line.strip() else line for line in lines]
    return "\n".join(fixed)


def normalize_text(text: str) -> str:
    # Strip null bytes (\x00) before anything else – PostgreSQL text columns
    # reject them with "invalid byte sequence for encoding UTF8: 0x00".
    cleaned = (text or "").replace("\x00", "")
    return " ".join(cleaned.split()).strip()


def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _looks_like_pdf(content: bytes) -> bool:
    return content.lstrip().startswith(b"%PDF-")


def _looks_like_html(content: bytes) -> bool:
    sample = content[:2048].decode("utf-8", errors="ignore").strip().lower()
    return "<html" in sample or "<!doctype html" in sample


def parse_html(content: bytes) -> tuple[str, str]:
    decoded = content.decode("utf-8", errors="replace")
    soup = BeautifulSoup(decoded, "html.parser")
    text = soup.get_text(separator=" ", strip=True)
    return decoded, text


def _parse_pdf_pdfplumber(content: bytes) -> str:
    """Extract text from a PDF using pdfplumber with RTL-aware word ordering.

    For each page pdfplumber returns every word with its (x, y) bounding box.
    Arabic PDFs often store words in visual right-to-left order in the content
    stream, so we cluster words by vertical band (y-position) and, for RTL
    lines, sort by x-position descending (rightmost word first) to restore
    natural reading order.

    Null bytes and Arabic presentation forms are normalised per _clean_word().
    """
    import pdfplumber  # optional dependency; imported lazily

    pages_text: list[str] = []
    with pdfplumber.open(BytesIO(content)) as pdf:
        for page in pdf.pages:
            words = page.extract_words(
                x_tolerance=5, y_tolerance=5, keep_blank_chars=False
            )
            if not words:
                pages_text.append("")
                continue

            # Group word tokens by approximate y-band (5pt buckets)
            line_map: dict[int, list[tuple[float, str]]] = {}
            for w in words:
                cleaned = _clean_word(w["text"])
                if not cleaned.strip():
                    continue
                y_bucket = round(w["top"] / 5) * 5
                line_map.setdefault(y_bucket, []).append((w["x0"], cleaned))

            lines: list[str] = []
            for y_bucket in sorted(line_map):
                tokens = line_map[y_bucket]
                line_preview = " ".join(t for _, t in tokens)
                if _has_rtl(line_preview):
                    # RTL line: sort words right-to-left (highest x first)
                    tokens.sort(key=lambda t: -t[0])
                else:
                    # LTR line: sort left-to-right (lowest x first)
                    tokens.sort(key=lambda t: t[0])
                lines.append(" ".join(t[::-1] if _has_rtl(t) else t for _, t in tokens))

            pages_text.append("\n".join(lines))

    return "\n\n".join(pages_text).strip()


def parse_pdf(content: bytes) -> str:
    """Extract text from a PDF, using pdfplumber for RTL/Arabic documents.

    Strategy:
      1. Attempt pdfplumber extraction with position-based RTL word ordering.
         This correctly reconstructs Arabic text by sorting words right-to-left
         within each line based on their x-coordinate in the page.
      2. If pdfplumber raises any exception (missing dependency, corrupt PDF,
         etc.), fall back to pypdf with NFKC normalisation.
    """
    try:
        text = _parse_pdf_pdfplumber(content)
        if text.strip():
            return text
        # pdfplumber returned empty — fall through to pypdf
    except Exception:
        pass

    # pypdf fallback: NFKC normalise (strips null bytes + presentation forms)
    reader = PdfReader(BytesIO(content))
    extracted = []
    for page in reader.pages:
        extracted.append(page.extract_text() or "")
    raw = "\n".join(extracted).strip()
    return _fix_rtl_text(raw)


def parse_plain_text(content: bytes) -> str:
    return content.decode("utf-8", errors="replace").strip()


def parse_docx(content: bytes) -> str:
    document = DocxDocument(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        value = (paragraph.text or "").strip()
        if value:
            parts.append(value)
    for table in document.tables:
        for row in table.rows:
            row_values = [(cell.text or "").strip() for cell in row.cells]
            merged = " | ".join([value for value in row_values if value])
            if merged:
                parts.append(merged)
    return "\n".join(parts).strip()


def _detect_delimiter(sample: str) -> str:
    """Guess the delimiter for a CSV-like file from its first few lines."""
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
        return dialect.delimiter
    except csv.Error:
        return ","


def parse_csv_like(content: bytes) -> str:
    """Parse CSV, TSV, DSV (pipe/semicolon) files into readable text.

    Each row is rendered as a pipe-separated line. The first row is
    treated as a header line.
    """
    text = content.decode("utf-8", errors="replace")
    # Strip null bytes early
    text = text.replace("\x00", "")
    if not text.strip():
        return ""

    delimiter = _detect_delimiter(text[:8192])
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    lines: list[str] = []
    for row in reader:
        cleaned = [cell.strip() for cell in row]
        if any(cleaned):
            lines.append(" | ".join(cleaned))
    return "\n".join(lines).strip()


def parse_excel(content: bytes) -> str:
    """Parse .xlsx / .xls files using openpyxl (xlsx) with xlrd fallback (xls)."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(BytesIO(content), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            if len(wb.sheetnames) > 1:
                parts.append(f"--- {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(cells):
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts).strip()
    except Exception:
        # Fallback: try xlrd for legacy .xls
        try:
            import xlrd

            wb = xlrd.open_workbook(file_contents=content)
            parts = []
            for sheet_idx in range(wb.nsheets):
                ws = wb.sheet_by_index(sheet_idx)
                if wb.nsheets > 1:
                    parts.append(f"--- {ws.name} ---")
                for row_idx in range(ws.nrows):
                    cells = [
                        str(ws.cell_value(row_idx, col)).strip()
                        for col in range(ws.ncols)
                    ]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts).strip()
        except Exception:
            raise


@dataclass
class ByteExtractionResult:
    status: str
    extraction_method: str
    extracted_text: Optional[str]
    normalized_text_hash: Optional[str]
    raw_html: Optional[str]
    ocr_provider_used: str = "none"
    fallback_stage: str = "none"
    warnings: list[str] | None = None
    error_code: Optional[str] = None


async def extract_text_from_bytes(
    *,
    content: bytes,
    content_type: str,
    file_name: str = "document",
    max_chars: Optional[int] = None,
) -> ByteExtractionResult:
    if len(content) > settings.extraction_max_bytes:
        return ByteExtractionResult(
            status="error",
            extraction_method="none",
            extracted_text=None,
            normalized_text_hash=None,
            raw_html=None,
            error_code="payload_too_large",
            warnings=[
                f"Uploaded payload exceeds {settings.extraction_max_bytes} bytes limit."
            ],
        )

    ctype = (content_type or "").lower()
    extension = Path(file_name or "").suffix.lower()
    warnings: list[str] = []
    raw_html: str | None = None
    parser_text = ""
    extraction_method = "parser_fallback"

    is_html = "text/html" in ctype or extension in (".html", ".htm")
    is_pdf = "application/pdf" in ctype or extension == ".pdf"
    is_text = ctype.startswith("text/") or extension in (".txt", ".md")
    is_docx = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        in ctype
        or extension == ".docx"
    )
    is_doc = "application/msword" in ctype or extension == ".doc"
    is_image = ctype.startswith("image/")
    is_csv_like = (
        "text/csv" in ctype
        or "text/tab-separated-values" in ctype
        or "text/dsv" in ctype
        or extension in (".csv", ".tsv", ".dsv")
    )
    is_excel = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" in ctype
        or "application/vnd.ms-excel" in ctype
        or extension in (".xlsx", ".xls")
    )
    is_rtf = "application/rtf" in ctype or "text/rtf" in ctype or extension == ".rtf"
    is_markdown = extension in (".md", ".markdown")

    if is_pdf and not _looks_like_pdf(content):
        if _looks_like_html(content):
            is_pdf = False
            is_html = True
            warnings.append(
                "Payload is labeled as PDF but appears to be HTML; using HTML parser."
            )
        else:
            is_pdf = False
            warnings.append(
                "Payload is labeled as PDF but does not contain a PDF header; skipping PDF parser/OCR."
            )

    if is_doc:
        return ByteExtractionResult(
            status="error",
            extraction_method="none",
            extracted_text=None,
            normalized_text_hash=None,
            raw_html=None,
            error_code="unsupported_file_type",
            warnings=[
                "Legacy .doc files are not supported yet. Upload .docx/PDF/text/image instead."
            ],
        )

    try:
        if is_html:
            raw_html, parser_text = parse_html(content)
            extraction_method = "parser_html"
        elif is_pdf:
            parser_text = parse_pdf(content)
            extraction_method = "parser_pdf"
        elif is_docx:
            parser_text = parse_docx(content)
            extraction_method = "parser_docx"
        elif is_csv_like:
            parser_text = parse_csv_like(content)
            extraction_method = "parser_csv"
        elif is_excel:
            parser_text = parse_excel(content)
            extraction_method = "parser_excel"
        elif is_rtf:
            # RTF: strip control words and extract plain text via a
            # lightweight approach; full RTF support would need striprtf.
            try:
                from striprtf.striprtf import rtf_to_text

                parser_text = rtf_to_text(content.decode("utf-8", errors="replace"))
            except ImportError:
                # Fallback: treat as plain text (some RTF files are readable)
                parser_text = parse_plain_text(content)
                warnings.append("striprtf not installed; RTF parsed as plain text.")
            extraction_method = "parser_rtf"
        elif is_text or is_markdown:
            parser_text = parse_plain_text(content)
            extraction_method = "parser_text"
        elif is_image:
            parser_text = ""
            extraction_method = "parser_image"
        else:
            parser_text = parse_plain_text(content)
            extraction_method = "parser_fallback"
    except Exception as exc:
        warnings.append(f"Parser failed: {exc}")
        parser_text = ""

    parsed_length = len(normalize_text(parser_text))
    needs_ocr = parsed_length < settings.ocr_min_text_chars and (is_pdf or is_image)

    # Arabic/RTL PDFs: always prefer OCR over parser extraction.
    # pdfplumber extracts Arabic chars in visual (LTR) glyph order with
    # word-boundary splitting artifacts.  Tesseract OCR reads the rendered
    # page and returns correct Arabic text with proper word boundaries.
    is_rtl_pdf = is_pdf and _has_rtl(parser_text)
    if is_rtl_pdf and not needs_ocr:
        needs_ocr = True

    ocr_provider_used = "none"
    fallback_stage = "none"
    final_text = parser_text

    if needs_ocr:
        primary_provider = create_ocr_provider(settings.ocr_primary_provider)
        primary = await primary_provider.extract_text(
            content=content,
            content_type=ctype or "application/octet-stream",
            file_name=file_name or "document",
        )
        if primary.ok and primary.text:
            final_text = primary.text
            extraction_method = "ocr_primary"
            ocr_provider_used = primary.provider
        else:
            warnings.append(primary.error or "Primary OCR provider failed.")
            fallback_stage = "secondary"
            secondary_provider = create_ocr_provider(settings.ocr_secondary_provider)
            secondary = await secondary_provider.extract_text(
                content=content,
                content_type=ctype or "application/octet-stream",
                file_name=file_name or "document",
            )
            if secondary.ok and secondary.text:
                final_text = secondary.text
                extraction_method = "ocr_secondary"
                ocr_provider_used = secondary.provider
            else:
                warnings.append(secondary.error or "Secondary OCR provider failed.")
                fallback_stage = "parser_only"
                extraction_method = f"{extraction_method}_fallback"
                ocr_provider_used = "none"

    max_chars_value = max_chars or settings.extraction_max_chars
    normalized = normalize_text(final_text)
    if len(normalized) > max_chars_value:
        normalized = normalized[:max_chars_value]
        warnings.append(f"Extracted text truncated to {max_chars_value} chars.")

    if not normalized:
        if settings.ocr_strict_mode:
            return ByteExtractionResult(
                status="error",
                extraction_method=extraction_method,
                extracted_text=None,
                normalized_text_hash=None,
                raw_html=raw_html,
                ocr_provider_used=ocr_provider_used,
                fallback_stage=fallback_stage,
                error_code="empty_extracted_text",
                warnings=warnings or ["No text could be extracted."],
            )
        warnings.append(
            "No text could be extracted. Returning empty content in non-strict mode."
        )

    return ByteExtractionResult(
        status="ok",
        extraction_method=extraction_method,
        extracted_text=normalized,
        normalized_text_hash=hash_text(normalized) if normalized else None,
        raw_html=raw_html,
        ocr_provider_used=ocr_provider_used,
        fallback_stage=fallback_stage,
        warnings=warnings,
    )
